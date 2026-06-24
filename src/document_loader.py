"""
document_loader.py
===================
Handles loading and text extraction from company knowledge-base documents.

Supports PDF, TXT and DOCX formats. Each loader returns plain text along
with simple metadata (source filename, page/section info where available)
so that downstream chunking and retrieval can cite sources accurately.
"""

from dataclasses import dataclass, field
from pathlib import Path
from typing import List

import docx  # python-docx
from pypdf import PdfReader

from src.utils import clean_text, get_logger, validate_file

logger = get_logger(__name__)


@dataclass
class LoadedDocument:
    """
    Represents a single loaded document (or a page/section within one),
    ready to be passed into the text splitter.

    Attributes:
        content: Extracted, cleaned text content.
        source: Original filename the content was extracted from.
        page: Optional page number (1-indexed) for PDFs. ``None`` for
            formats without a native page concept (e.g. TXT, DOCX).
        metadata: Additional free-form metadata dictionary.
    """

    content: str
    source: str
    page: int = None
    metadata: dict = field(default_factory=dict)


class DocumentLoaderError(Exception):
    """Raised when a document cannot be loaded or parsed."""


class DocumentLoader:
    """
    Loads and extracts text from PDF, TXT and DOCX files.

    This class follows a strategy-style design: a single public entry point
    (`load`) dispatches to the appropriate private extraction method based
    on file extension, keeping the public API simple while remaining easy
    to extend with new formats.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}

    def load(self, file_path: str) -> List[LoadedDocument]:
        """
        Load a single file and return its extracted content as a list of
        ``LoadedDocument`` objects (one per page for PDFs, one entry for
        TXT/DOCX).

        Args:
            file_path: Absolute or relative path to the file on disk.

        Returns:
            List of ``LoadedDocument`` instances.

        Raises:
            DocumentLoaderError: If the file is missing, unsupported, or
                fails to parse.
        """
        path = Path(file_path)
        if not path.exists():
            raise DocumentLoaderError(f"File not found: {file_path}")

        # Validate extension & size before attempting to parse.
        try:
            validate_file(path.name, path.stat().st_size)
        except Exception as exc:
            raise DocumentLoaderError(str(exc)) from exc

        extension = path.suffix.lower()
        logger.info("Loading document '%s' (%s)", path.name, extension)

        try:
            if extension == ".pdf":
                return self._load_pdf(path)
            elif extension == ".txt":
                return self._load_txt(path)
            elif extension == ".docx":
                return self._load_docx(path)
            else:
                raise DocumentLoaderError(
                    f"Unsupported file extension '{extension}'. "
                    f"Supported: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
                )
        except DocumentLoaderError:
            raise
        except Exception as exc:
            logger.error("Failed to parse '%s': %s", path.name, exc)
            raise DocumentLoaderError(
                f"Failed to parse '{path.name}': {exc}"
            ) from exc

    def load_directory(self, directory_path: str) -> List[LoadedDocument]:
        """
        Load every supported file found in a directory (non-recursive).

        Args:
            directory_path: Path to a directory containing knowledge-base
                files.

        Returns:
            Combined list of ``LoadedDocument`` objects from all files.
        """
        directory = Path(directory_path)
        if not directory.exists() or not directory.is_dir():
            logger.warning("Directory '%s' does not exist.", directory_path)
            return []

        documents: List[LoadedDocument] = []
        for file_path in sorted(directory.iterdir()):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    documents.extend(self.load(str(file_path)))
                except DocumentLoaderError as exc:
                    logger.error("Skipping '%s': %s", file_path.name, exc)
        return documents

    # ------------------------------------------------------------------ #
    # Private format-specific extraction methods
    # ------------------------------------------------------------------ #
    def _load_pdf(self, path: Path) -> List[LoadedDocument]:
        """Extract text from a PDF, page by page."""
        documents: List[LoadedDocument] = []
        try:
            reader = PdfReader(str(path))
        except Exception as exc:
            raise DocumentLoaderError(
                f"'{path.name}' could not be read as a valid PDF: {exc}"
            ) from exc

        if len(reader.pages) == 0:
            raise DocumentLoaderError(f"'{path.name}' contains no pages.")

        for page_num, page in enumerate(reader.pages, start=1):
            raw_text = page.extract_text() or ""
            text = clean_text(raw_text)
            if text:
                documents.append(
                    LoadedDocument(
                        content=text,
                        source=path.name,
                        page=page_num,
                        metadata={"file_type": "pdf", "total_pages": len(reader.pages)},
                    )
                )

        if not documents:
            logger.warning(
                "'%s' produced no extractable text (likely a scanned/image PDF).",
                path.name,
            )
        return documents

    def _load_txt(self, path: Path) -> List[LoadedDocument]:
        """Extract text from a plain-text file."""
        try:
            raw_text = path.read_text(encoding="utf-8", errors="replace")
        except Exception as exc:
            raise DocumentLoaderError(
                f"'{path.name}' could not be read as text: {exc}"
            ) from exc

        text = clean_text(raw_text)
        if not text:
            raise DocumentLoaderError(f"'{path.name}' is empty.")

        return [
            LoadedDocument(
                content=text,
                source=path.name,
                page=None,
                metadata={"file_type": "txt"},
            )
        ]

    def _load_docx(self, path: Path) -> List[LoadedDocument]:
        """Extract text from a Word (.docx) document, including tables."""
        try:
            document = docx.Document(str(path))
        except Exception as exc:
            raise DocumentLoaderError(
                f"'{path.name}' could not be read as a valid DOCX file: {exc}"
            ) from exc

        parts: List[str] = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    parts.append(row_text)

        text = clean_text("\n".join(parts))
        if not text:
            raise DocumentLoaderError(f"'{path.name}' contains no extractable text.")

        return [
            LoadedDocument(
                content=text,
                source=path.name,
                page=None,
                metadata={"file_type": "docx"},
            )
        ]
